from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc = Document()

    # Set default font to Arial 11
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # 1. Header (Highlighted)
    p = doc.add_paragraph()
    runner = p.add_run('SPEAK UP SURVEY:          SIATech Network Consolidated Reports')
    runner.font.bold = True
    runner.font.highlight_color = 7 # Yellow highlight

    doc.add_paragraph() # spacer

    # 2. NETWORK Data (labeled as District)
    p = doc.add_paragraph()
    r = p.add_run('NETWORK Data (labeled as District)')
    r.font.bold = True
    
    doc.add_paragraph('Students             Grades 9-12')
    doc.add_paragraph('Teachers')
    doc.add_paragraph('School Site Administrators')
    doc.add_paragraph('District Administrators')
    doc.add_paragraph('Tech Leaders')
    doc.add_paragraph('Parents')
    doc.add_paragraph('Community Members')
    
    doc.add_paragraph() # spacer

    # 3. Individual School Results
    p = doc.add_paragraph()
    r = p.add_run('Individual School Results')
    r.font.bold = True
    r.font.size = Pt(12)

    p = doc.add_paragraph()
    r = p.add_run('CALIFORNIA')
    r.font.bold = True

    doc.add_paragraph('SIATech Sacramento             Job Corps')
    doc.add_paragraph('SIATech San Jose                 Job Corps')
    doc.add_paragraph('SIATech Treasure Island         Job Corps')
    doc.add_paragraph('SIATech Long Beach             Job Corps')
    doc.add_paragraph('SIATech Los Angeles            Job Corps')
    doc.add_paragraph('SIATech Inland Empire          Job Corps')
    doc.add_paragraph('SIATech San Diego                Job Corps')
    
    doc.add_paragraph() # spacer
    
    doc.add_paragraph('SIATech Moreno Valley  Riverside        Independent Study')
    doc.add_paragraph('SIATech Indio               Riverside        Independent Study')
    doc.add_paragraph('SIATech Perris              Riverside        Independent Study')
    doc.add_paragraph('SIATech South Bay        San Diego      Independent Study')
    
    doc.add_paragraph() # spacer
    
    doc.add_paragraph('SAS      Pico Union       Los Angeles    Independent Study')
    doc.add_paragraph('SAS      Boyle Heights   Los Angeles    Independent Study')
    
    doc.add_paragraph() # spacer

    # 4. ARKANSAS
    p = doc.add_paragraph()
    r = p.add_run('ARKANSAS')
    r.font.bold = True
    
    doc.add_paragraph('SIATech Little Rock*      Little Rock      Community')

    doc.add_paragraph() # spacer

    # 5. FLORIDA
    p = doc.add_paragraph()
    r = p.add_run('FLORIDA')
    r.font.bold = True
    
    doc.add_paragraph('MYcroSchool Gainesville*    Alachua       Community')
    doc.add_paragraph('MYcroSchool Jacksonville*  Duval           Community')
    doc.add_paragraph('MYcroSchool Pinellas*         Pinellas        Community')
    doc.add_paragraph('SIATech Gainesville*           Alachua        Community')

    doc.add_paragraph() # spacer

    # 6. Acronyms
    p = doc.add_paragraph()
    r = p.add_run('Acronyms/ *')
    r.font.bold = True
    r.font.underline = True

    p = doc.add_paragraph()
    r = p.add_run('SIATech, Inc. (School for Integrated Academics and Technologies, Inc.)')
    r.font.italic = True
    
    p = doc.add_paragraph()
    r = p.add_run('SAS (SIATech Academy South, Inc.)')
    r.font.italic = True
    
    p = doc.add_paragraph()
    r = p.add_run('*Local Education Agency (LEA)')
    r.font.italic = True

    doc.add_paragraph() # spacer
    doc.add_paragraph() # spacer
    doc.add_paragraph() # spacer
    doc.add_paragraph() # spacer

    # Footer
    p = doc.add_paragraph()
    r = p.add_run('SIATech, Inc. OY Career & College Pathways Program                   Appendix G  Page 538')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save
    doc.save('SPEAK_UP_SURVEY_Transcription.docx')

if __name__ == '__main__':
    main()
