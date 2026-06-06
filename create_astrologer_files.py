import os
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

# Define the exact text content
text_content = """7. INTELLECTUAL PROPERTY
Any content created by the Astrologer specifically for the Astro Soul Path Platform (written reports, recorded sessions with customer consent, customised horoscopes, video content) shall be jointly owned by the Astrologer and Astro Soul Path Private Limited unless otherwise agreed in writing.
The Astrologer retains ownership of their general astrological methodology, teaching materials, and content created independently outside the Platform. Astro Soul Path Private Limited retains ownership of its technology, brand, trademarks, and Platform infrastructure.

8. CONFIDENTIALITY & DATA PROTECTION
The Astrologer agrees to keep confidential all customer personal data, session details, business information, and Platform operational data disclosed during the term of this Agreement and for 3 years thereafter.
The Astrologer shall comply with the Information Technology Act, 2000, and applicable data protection norms, including the Digital Personal Data Protection Act, 2023. Customer data shall not be stored locally on the Astrologer's personal devices beyond session requirements.

9. TERM & TERMINATION
a) This Agreement is valid for 12 months from the Agreement Date and shall auto-renew unless either party provides 30 days' written notice of non-renewal.
b) Either party may terminate this Agreement with 30 days' written notice without cause.
c) Astro Soul Path Private Limited may terminate immediately (without notice) in cases of: fraud, misrepresentation of credentials, customer harassment, data breach, or material breach of this Agreement.
d) Upon termination, any pending verified earnings shall be paid within 15 business days.

10. DISPUTE RESOLUTION
Any dispute arising out of or relating to this Agreement shall first be resolved through good-faith negotiation within 15 days of written notice. If unresolved, disputes shall be referred to arbitration under the Arbitration and Conciliation Act, 1996. Arbitration shall be conducted in Delhi, India, in the English language.

11. REPRESENTATIONS & WARRANTIES
The Astrologer represents and warrants that:
• They are of legal age (18 years or above) and have the legal capacity to enter this Agreement.
• All credentials, qualifications, and experience stated in their profile are accurate and verifiable.
• They are not bound by any existing agreement that would conflict with this Agreement.
• They hold a valid bank account in their own name for payment purposes.
• They will obtain and maintain all necessary registrations (GST, professional licences) as applicable."""

def generate_txt():
    with open("Astrologer_Partnership_Agreement.txt", "w", encoding="utf-8") as f:
        f.write(text_content)
    print("Astrologer_Partnership_Agreement.txt created.")

def generate_docx():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    lines = text_content.split('\n')
    for line in lines:
        if line.strip():
            p = doc.add_paragraph()
            # If line is a section heading, make it bold
            if (line.startswith("7.") or line.startswith("8.") or 
                line.startswith("9.") or line.startswith("10.") or 
                line.startswith("11.")):
                r = p.add_run(line)
                r.font.bold = True
            else:
                p.add_run(line)
        else:
            doc.add_paragraph() # spacing
            
    doc.save("Astrologer_Partnership_Agreement.docx")
    print("Astrologer_Partnership_Agreement.docx created.")

def generate_pdf():
    filename = "Astrologer_Partnership_Agreement.pdf"
    c = canvas.Canvas(filename, pagesize=letter)
    width, height = letter
    
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, height - 50, "Astrologer Partnership Agreement")
    c.line(50, height - 55, width - 50, height - 55)
    
    c.setFont("Helvetica", 10)
    
    y = height - 80
    lines = text_content.split('\n')
    for line in lines:
        if y < 60:
            c.showPage()
            c.setFont("Helvetica", 10)
            y = height - 50
            
        if line.strip():
            # Bold for headings
            if (line.startswith("7.") or line.startswith("8.") or 
                line.startswith("9.") or line.startswith("10.") or 
                line.startswith("11.")):
                c.setFont("Helvetica-Bold", 10)
            else:
                c.setFont("Helvetica", 10)
            
            # Simple text wrap for lines longer than 85 chars
            words = line.split()
            current_line = ""
            for word in words:
                if len(current_line + " " + word) < 85:
                    current_line += (" " if current_line else "") + word
                else:
                    c.drawString(50, y, current_line)
                    y -= 14
                    current_line = word
            if current_line:
                c.drawString(50, y, current_line)
                y -= 14
        else:
            y -= 8
            
    c.save()
    print("Astrologer_Partnership_Agreement.pdf created.")

if __name__ == "__main__":
    generate_txt()
    generate_docx()
    generate_pdf()
